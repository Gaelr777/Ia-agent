import docx
import pytest

from pipeline.cv_classifier.extract_text import UnsupportedFileType, extract_text


def test_extract_text_rejects_unsupported_extension(tmp_path):
    path = tmp_path / "cv.txt"
    path.write_text("hola")
    with pytest.raises(UnsupportedFileType):
        extract_text(path)


def test_extract_text_docx_reads_paragraphs_and_tables(tmp_path):
    document = docx.Document()
    document.add_paragraph("Ingeniero de producción con 5 años de experiencia.")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Habilidad"
    table.rows[0].cells[1].text = "Excel"
    path = tmp_path / "cv.docx"
    document.save(path)

    text = extract_text(path)

    assert "Ingeniero de producción" in text
    assert "Habilidad | Excel" in text
