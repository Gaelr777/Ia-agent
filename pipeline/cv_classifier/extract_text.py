"""Extracción de texto plano desde CV en PDF o DOCX."""
from pathlib import Path

import docx
import pdfplumber


class UnsupportedFileType(ValueError):
    pass


def extract_text(file_path: str | Path) -> str:
    path = Path(file_path)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix == ".docx":
        return _extract_docx(path)
    raise UnsupportedFileType(
        f"Tipo de archivo no soportado: {suffix}. Solo .pdf y .docx."
    )


def _extract_pdf(path: Path) -> str:
    chunks = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            chunks.append(text)
    return "\n".join(chunks).strip()


def _extract_docx(path: Path) -> str:
    document = docx.Document(str(path))
    paragraphs = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            paragraphs.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(paragraphs).strip()
